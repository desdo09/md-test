"""
create_sample_docs.py
---------------------
Generates sample test documents in the input/ folder so you can run
convert.py without needing any real documents.

Run once:  python create_sample_docs.py
"""

from pathlib import Path

INPUT_DIR = Path("input")
INPUT_DIR.mkdir(exist_ok=True)


# ── 1. Plain HTML document ───────────────────────────────────────────────────

html_content = """\
<!DOCTYPE html>
<html lang="en">
<head><meta charset="UTF-8"><title>Sample Report</title></head>
<body>
  <h1>Quarterly Sales Report</h1>
  <p>This is a <strong>sample HTML document</strong> used to test MarkItDown conversion.</p>

  <h2>Summary</h2>
  <p>Overall performance this quarter was <em>above expectations</em>.</p>

  <h2>Key Metrics</h2>
  <table border="1" cellpadding="4">
    <thead>
      <tr><th>Region</th><th>Revenue ($)</th><th>Growth</th></tr>
    </thead>
    <tbody>
      <tr><td>North</td><td>1,200,000</td><td>+12%</td></tr>
      <tr><td>South</td><td>980,000</td><td>+8%</td></tr>
      <tr><td>East</td><td>1,450,000</td><td>+18%</td></tr>
      <tr><td>West</td><td>870,000</td><td>+5%</td></tr>
    </tbody>
  </table>

  <h2>Action Items</h2>
  <ul>
    <li>Increase marketing spend in the West region</li>
    <li>Launch new product line in Q3</li>
    <li>Review pricing strategy for South region</li>
  </ul>

  <blockquote>
    <p>"Execution is everything." – Internal Memo</p>
  </blockquote>

  <h2>Code Sample</h2>
  <pre><code>def calculate_growth(prev, curr):
    return (curr - prev) / prev * 100
  </code></pre>
</body>
</html>
"""

(INPUT_DIR / "sample_report.html").write_text(html_content, encoding="utf-8")
print("Created: input/sample_report.html")


# ── 2. CSV spreadsheet ────────────────────────────────────────────────────────

csv_content = """\
Name,Department,Role,Salary,Start Date
Alice Johnson,Engineering,Senior Engineer,95000,2020-03-15
Bob Smith,Marketing,Marketing Manager,82000,2019-07-01
Carol White,Engineering,Tech Lead,110000,2018-01-10
David Brown,Sales,Account Executive,75000,2021-11-20
Eve Davis,HR,HR Specialist,68000,2022-06-05
Frank Wilson,Engineering,Junior Engineer,72000,2023-02-14
"""

(INPUT_DIR / "employees.csv").write_text(csv_content, encoding="utf-8")
print("Created: input/employees.csv")


# ── 3. JSON document ─────────────────────────────────────────────────────────

import json

json_content = {
    "project": "Alpha Initiative",
    "status": "In Progress",
    "owner": "Engineering Team",
    "milestones": [
        {"name": "Design Phase",    "due": "2026-02-28", "status": "Completed"},
        {"name": "Development",     "due": "2026-04-15", "status": "In Progress"},
        {"name": "QA & Testing",    "due": "2026-05-01", "status": "Pending"},
        {"name": "Go-Live",         "due": "2026-05-15", "status": "Pending"},
    ],
    "risks": [
        "Resource availability in March",
        "Third-party API integration delays",
    ],
    "notes": "All stakeholders have been briefed. Weekly syncs every Monday at 10 AM."
}

(INPUT_DIR / "project_status.json").write_text(
    json.dumps(json_content, indent=2), encoding="utf-8"
)
print("Created: input/project_status.json")


# ── 4. Word document (.docx) — requires python-docx ──────────────────────────

try:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Technical Specification", level=1)

    doc.add_paragraph(
        "This document outlines the technical requirements for the new data pipeline."
    )

    doc.add_heading("1. Overview", level=2)
    doc.add_paragraph(
        "The pipeline ingests raw event data from multiple sources, transforms it, "
        "and loads it into the data warehouse on an hourly schedule."
    )

    doc.add_heading("2. Architecture", level=2)
    doc.add_paragraph("Components involved:")
    for item in ["Apache Kafka (ingestion)", "Apache Spark (transformation)", "Snowflake (storage)"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. SLAs", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Metric", "Target", "Current"
    for metric, target, current in [
        ("Latency",     "< 5 min",  "3.2 min"),
        ("Throughput",  "> 10K/s",  "12.4K/s"),
        ("Error rate",  "< 0.1%",   "0.04%"),
    ]:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = metric, target, current

    doc.add_heading("4. Next Steps", level=2)
    doc.add_paragraph("Review and sign off on this spec by end of sprint.")

    doc.save(INPUT_DIR / "tech_spec.docx")
    print("Created: input/tech_spec.docx")

except ImportError:
    print("Skipped .docx sample (python-docx not installed yet — run after pip install).")


print("\nDone. Run:  python convert.py")
